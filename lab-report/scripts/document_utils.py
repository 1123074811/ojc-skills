"""
综合文档处理模块
集成 pdf skill 和 xlsx skill 功能
提供统一的文档处理接口
"""

import os
import sys
from typing import List, Optional, Dict, Any
from pathlib import Path

# 添加脚本目录到路径
script_dir = os.path.dirname(os.path.abspath(__file__))
if script_dir not in sys.path:
    sys.path.insert(0, script_dir)

from docx_handler import LabReportDocument


class DocumentProcessor:
    """
    综合文档处理器
    统一处理 DOCX、PDF、XLSX 文件
    """
    
    def __init__(self, work_dir: str = None):
        """
        初始化文档处理器
        
        Args:
            work_dir: 工作目录（用于临时文件）
        """
        self.work_dir = work_dir or os.getcwd()
        self._docx_available = False
        self._pdf_available = False
        self._xlsx_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查可用的依赖库"""
        # 检查 DOCX
        try:
            import docx
            self._docx_available = True
            print("[DOCX] python-docx 可用")
        except ImportError:
            print("[DOCX] python-docx 未安装")
        
        # 检查 PDF
        try:
            import pypdf
            self._pdf_available = True
            print("[PDF] pypdf 可用")
        except ImportError:
            try:
                import PyPDF2
                self._pdf_available = True
                print("[PDF] PyPDF2 可用")
            except ImportError:
                print("[PDF] PDF 库未安装")
        
        # 检查 XLSX
        try:
            import openpyxl
            self._xlsx_available = True
            print("[XLSX] openpyxl 可用")
        except ImportError:
            print("[XLSX] openpyxl 未安装")
    
    # ==================== PDF 处理功能 ====================
    
    def merge_pdfs(self, pdf_files: List[str], output_path: str) -> bool:
        """
        合并多个 PDF 文件
        
        Args:
            pdf_files: PDF 文件路径列表
            output_path: 输出文件路径
        
        Returns:
            bool: 是否成功
        """
        if not self._pdf_available:
            print("错误: PDF 库未安装，无法合并 PDF")
            return False
        
        print(f"正在合并 {len(pdf_files)} 个 PDF 文件...")
        
        try:
            # 优先使用 pypdf
            try:
                from pypdf import PdfMerger
                merger = PdfMerger()
            except ImportError:
                from PyPDF2 import PdfMerger
                merger = PdfMerger()
            
            for pdf_file in pdf_files:
                if os.path.exists(pdf_file):
                    merger.append(pdf_file)
                    print(f"  已添加: {os.path.basename(pdf_file)}")
                else:
                    print(f"  警告: 文件不存在: {pdf_file}")
            
            merger.write(output_path)
            merger.close()
            
            print(f"PDF 合并完成: {output_path}")
            return True
            
        except Exception as e:
            print(f"合并 PDF 时出错: {e}")
            return False
    
    def extract_pdf_text(self, pdf_path: str, output_txt: str = None) -> str:
        """
        提取 PDF 文本内容
        
        Args:
            pdf_path: PDF 文件路径
            output_txt: 输出文本文件路径（可选）
        
        Returns:
            str: 提取的文本内容
        """
        if not self._pdf_available:
            print("错误: PDF 库未安装，无法提取文本")
            return ""
        
        print(f"正在提取 PDF 文本: {pdf_path}")
        
        try:
            # 优先使用 pypdf
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader
            
            reader = PdfReader(pdf_path)
            text = ""
            
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += f"\n--- 第 {i+1} 页 ---\n"
                text += page_text
            
            # 保存到文件
            if output_txt:
                with open(output_txt, 'w', encoding='utf-8') as f:
                    f.write(text)
                print(f"文本已保存: {output_txt}")
            
            return text
            
        except Exception as e:
            print(f"提取 PDF 文本时出错: {e}")
            return ""
    
    def split_pdf(self, pdf_path: str, output_dir: str, 
                  pages_per_file: int = 1) -> List[str]:
        """
        拆分 PDF 文件
        
        Args:
            pdf_path: PDF 文件路径
            output_dir: 输出目录
            pages_per_file: 每个文件的页数
        
        Returns:
            List[str]: 生成的文件路径列表
        """
        if not self._pdf_available:
            print("错误: PDF 库未安装，无法拆分 PDF")
            return []
        
        print(f"正在拆分 PDF: {pdf_path}")
        
        try:
            # 优先使用 pypdf
            try:
                from pypdf import PdfReader, PdfWriter
            except ImportError:
                from PyPDF2 import PdfReader, PdfWriter
            
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            
            os.makedirs(output_dir, exist_ok=True)
            
            output_files = []
            file_count = 0
            
            for i in range(0, total_pages, pages_per_file):
                writer = PdfWriter()
                
                end_page = min(i + pages_per_file, total_pages)
                for page_num in range(i, end_page):
                    writer.add_page(reader.pages[page_num])
                
                output_filename = f"split_{file_count+1:03d}.pdf"
                output_path = os.path.join(output_dir, output_filename)
                
                with open(output_path, 'wb') as output_file:
                    writer.write(output_file)
                
                output_files.append(output_path)
                print(f"  生成: {output_filename} (页 {i+1}-{end_page})")
                
                file_count += 1
            
            print(f"PDF 拆分完成，共 {file_count} 个文件")
            return output_files
            
        except Exception as e:
            print(f"拆分 PDF 时出错: {e}")
            return []
    
    # ==================== XLSX 处理功能 ====================
    
    def create_excel_from_table(self, table_data: List[List[Any]], 
                                headers: List[str],
                                output_path: str,
                                sheet_name: str = "Sheet1") -> bool:
        """
        从表格数据创建 Excel 文件
        
        Args:
            table_data: 表格数据（二维列表）
            headers: 表头列表
            output_path: 输出文件路径
            sheet_name: 工作表名称
        
        Returns:
            bool: 是否成功
        """
        if not self._xlsx_available:
            print("错误: openpyxl 未安装，无法创建 Excel")
            return False
        
        print(f"正在创建 Excel 文件: {output_path}")
        
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            
            # 写入表头
            ws.append(headers)
            
            # 设置表头样式
            header_font = Font(bold=True, size=11)
            for cell in ws[1]:
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # 写入数据
            for row_data in table_data:
                ws.append(row_data)
            
            # 自动调整列宽
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            wb.save(output_path)
            print(f"Excel 文件创建完成: {output_path}")
            return True
            
        except Exception as e:
            print(f"创建 Excel 时出错: {e}")
            return False
    
    def export_docx_tables_to_excel(self, docx_path: str, 
                                    output_dir: str = None) -> List[str]:
        """
        导出 DOCX 中的所有表格到 Excel 文件
        
        Args:
            docx_path: DOCX 文件路径
            output_dir: 输出目录（可选）
        
        Returns:
            List[str]: 生成的 Excel 文件路径列表
        """
        if not self._xlsx_available:
            print("错误: openpyxl 未安装，无法导出表格")
            return []
        
        print(f"正在导出 DOCX 表格: {docx_path}")
        
        try:
            from docx import Document
            
            doc = Document(docx_path)
            
            if not doc.tables:
                print("文档中没有表格")
                return []
            
            if output_dir is None:
                output_dir = os.path.dirname(docx_path) or self.work_dir
            
            os.makedirs(output_dir, exist_ok=True)
            
            output_files = []
            base_name = Path(docx_path).stem
            
            for i, table in enumerate(doc.tables):
                # 提取表格数据
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    headers = table_data[0]
                    data = table_data[1:] if len(table_data) > 1 else []
                    
                    output_path = os.path.join(
                        output_dir, 
                        f"{base_name}_table_{i+1}.xlsx"
                    )
                    
                    if self.create_excel_from_table(data, headers, output_path):
                        output_files.append(output_path)
            
            print(f"共导出 {len(output_files)} 个表格")
            return output_files
            
        except Exception as e:
            print(f"导出表格时出错: {e}")
            return []
    
    def create_test_case_excel(self, test_cases: List[Dict[str, str]], 
                               output_path: str) -> bool:
        """
        创建测试用例 Excel 文件
        
        Args:
            test_cases: 测试用例列表，每项为字典
            output_path: 输出文件路径
        
        Returns:
            bool: 是否成功
        """
        if not test_cases:
            print("错误: 测试用例为空")
            return False
        
        # 获取所有字段
        headers = list(test_cases[0].keys())
        
        # 转换为二维列表
        data = []
        for tc in test_cases:
            row = [tc.get(h, "") for h in headers]
            data.append(row)
        
        return self.create_excel_from_table(data, headers, output_path, "测试用例")
    
    # ==================== 综合功能 ====================
    
    def convert_docx_to_pdf(self, docx_path: str, 
                           output_pdf: str = None) -> bool:
        """
        将 DOCX 转换为 PDF
        
        注意：此功能需要 LibreOffice 或 Microsoft Word
        
        Args:
            docx_path: DOCX 文件路径
            output_pdf: 输出 PDF 路径（可选）
        
        Returns:
            bool: 是否成功
        """
        if output_pdf is None:
            output_pdf = docx_path.replace('.docx', '.pdf')
        
        print(f"正在转换 DOCX 到 PDF: {docx_path}")
        
        # 尝试使用 LibreOffice
        try:
            import subprocess
            
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(output_pdf) or '.',
                docx_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                print(f"转换成功: {output_pdf}")
                return True
            else:
                print(f"LibreOffice 转换失败: {result.stderr}")
                return False
                
        except FileNotFoundError:
            print("错误: 未找到 LibreOffice，无法转换")
            print("请安装 LibreOffice 或使用其他转换工具")
            return False
        except Exception as e:
            print(f"转换时出错: {e}")
            return False


def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='综合文档处理工具')
    parser.add_argument('command', choices=[
        'merge-pdfs', 'extract-pdf', 'split-pdf',
        'export-tables', 'create-excel'
    ], help='命令')
    parser.add_argument('--input', '-i', nargs='+', required=True, 
                       help='输入文件')
    parser.add_argument('--output', '-o', required=True, 
                       help='输出文件或目录')
    
    args = parser.parse_args()
    
    processor = DocumentProcessor()
    
    if args.command == 'merge-pdfs':
        success = processor.merge_pdfs(args.input, args.output)
    elif args.command == 'extract-pdf':
        text = processor.extract_pdf_text(args.input[0], args.output)
        success = bool(text)
    elif args.command == 'split-pdf':
        files = processor.split_pdf(args.input[0], args.output)
        success = bool(files)
    elif args.command == 'export-tables':
        files = processor.export_docx_tables_to_excel(args.input[0], args.output)
        success = bool(files)
    else:
        print(f"未知命令: {args.command}")
        success = False
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
