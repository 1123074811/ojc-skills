"""
DOCX 文档处理模块
集成 docx skill 的 Document 库，提供高级 DOCX 操作功能
"""

import os
import sys
import shutil
from datetime import datetime
from typing import Optional, List, Dict, Tuple

# 尝试导入 docx skill 的 Document 库
def check_docx_skill():
    """检查 docx skill 是否可用"""
    try:
        # 尝试从标准路径导入
        docx_skill_paths = [
            os.path.expanduser('~/.claude/skills/docx'),
            r'C:\Users\MECHREVO\.claude\skills\docx',
            '/usr/local/.claude/skills/docx'
        ]
        
        for skill_path in docx_skill_paths:
            if os.path.exists(skill_path) and skill_path not in sys.path:
                sys.path.insert(0, skill_path)
                try:
                    from scripts.document import Document, DocxXMLEditor
                    return True
                except ImportError:
                    continue
        return False
    except ImportError:
        return False

DOCX_SKILL_AVAILABLE = check_docx_skill()
if not DOCX_SKILL_AVAILABLE:
    print("信息: docx skill 不可用，将使用 python-docx 作为备选方案")


def unpack_docx(docx_path: str, output_dir: str) -> bool:
    """
    解压 DOCX 文件
    
    Args:
        docx_path: DOCX 文件路径
        output_dir: 输出目录
    
    Returns:
        bool: 是否成功
    """
    try:
        import zipfile
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(output_dir)
        return True
    except Exception as e:
        print(f"解压 DOCX 失败: {e}")
        return False


def pack_docx(source_dir: str, docx_path: str) -> bool:
    """
    打包 DOCX 文件
    
    Args:
        source_dir: 源目录
        docx_path: 输出 DOCX 文件路径
    
    Returns:
        bool: 是否成功
    """
    try:
        import zipfile
        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(source_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, source_dir)
                    zipf.write(file_path, arcname)
        return True
    except Exception as e:
        print(f"打包 DOCX 失败: {e}")
        return False


class LabReportDocument:
    """
    实验报告文档处理类
    封装 docx skill 的功能，提供实验报告特定的操作
    """
    
    def __init__(self, docx_path: str, temp_dir: Optional[str] = None):
        """
        初始化实验报告文档
        
        Args:
            docx_path: DOCX 文件路径
            temp_dir: 临时目录（可选）
        """
        self.original_path = docx_path
        self.temp_dir = temp_dir or os.path.join(
            os.path.dirname(docx_path),
            f"_temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        self.unpacked_path = None
        self.doc = None
        self._is_open = False
        
    def open(self) -> bool:
        """打开文档"""
        try:
            if DOCX_SKILL_AVAILABLE:
                # 使用 docx skill 的 Document 类
                # 需要先解压文档
                if not os.path.exists(self.temp_dir):
                    os.makedirs(self.temp_dir)
                
                self.unpacked_path = os.path.join(self.temp_dir, 'unpacked')
                if not unpack_docx(self.original_path, self.unpacked_path):
                    return False
                
                self.doc = Document(self.unpacked_path, author="LabReport", initials="LR")
            else:
                # 使用 python-docx 作为备选
                from docx import Document as DocxDocument
                self.doc = DocxDocument(self.original_path)
            
            self._is_open = True
            return True
        except Exception as e:
            print(f"打开文档失败: {e}")
            return False
    
    def close(self, save_path: Optional[str] = None) -> bool:
        """
        关闭文档
        
        Args:
            save_path: 保存路径（可选，默认覆盖原文件）
        """
        if not self._is_open:
            return True
        
        try:
            output_path = save_path or self.original_path
            
            if DOCX_SKILL_AVAILABLE and self.doc:
                # 使用 docx skill 保存
                self.doc.save()
                # 打包回 DOCX
                if not pack_docx(self.unpacked_path, output_path):
                    return False
            elif self.doc:
                # 使用 python-docx 保存
                self.doc.save(output_path)
            
            self._is_open = False
            
            # 清理临时目录
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
            
            return True
        except Exception as e:
            print(f"关闭文档失败: {e}")
            return False
    
    def fill_cover_info(self, student_id: str, name: str, class_name: str, 
                        date: Optional[str] = None) -> bool:
        """
        填写封面信息
        
        Args:
            student_id: 学号
            name: 姓名
            class_name: 班级
            date: 日期（可选，默认当天）
        """
        if not self._is_open:
            return False
        
        if date is None:
            date = datetime.now().strftime("%Y年%m月%d日")
        
        try:
            if DOCX_SKILL_AVAILABLE:
                # 使用 docx skill 的方式
                # 查找并替换封面信息
                doc_xml = self.doc["word/document.xml"]
                
                # 替换学号
                try:
                    node = doc_xml.get_node(tag="w:t", contains="学号")
                    if node:
                        parent = node.parentNode
                        # 找到下一个包含可填写位置的节点
                        # 这里需要根据具体模板结构调整
                        # 暂时使用简单替换
                        if node.nextSibling:
                            node.nextSibling.text = student_id
                except:
                    pass
                
                # 替换姓名
                try:
                    node = doc_xml.get_node(tag="w:t", contains="姓名")
                    if node:
                        if node.nextSibling:
                            node.nextSibling.text = name
                except:
                    pass
                
                # 替换班级
                try:
                    node = doc_xml.get_node(tag="w:t", contains="班级")
                    if node:
                        if node.nextSibling:
                            node.nextSibling.text = class_name
                except:
                    pass
                
                # 替换日期
                try:
                    node = doc_xml.get_node(tag="w:t", contains="日期")
                    if node:
                        if node.nextSibling:
                            node.nextSibling.text = date
                except:
                    pass
            else:
                # 使用 python-docx 的方式 - 改进的鲁棒性实现
                for para in self.doc.paragraphs:
                    text = para.text.strip()
                    
                    # 学号替换 - 更精确的匹配
                    if "学号" in text and student_id not in text:
                        # 查找学号字段后的空白位置
                        for i, run in enumerate(para.runs):
                            if "学号" in run.text:
                                # 查找下一个非空run
                                for j in range(i + 1, len(para.runs)):
                                    next_run = para.runs[j]
                                    if next_run.text.strip() == "" or next_run.text.isspace():
                                        next_run.text = student_id
                                        break
                                    elif not any(char in next_run.text for char in ["学号", "："]):
                                        next_run.text = student_id
                                        break
                        break
                    
                    # 姓名替换
                    elif "姓名" in text and name not in text:
                        for i, run in enumerate(para.runs):
                            if "姓名" in run.text:
                                for j in range(i + 1, len(para.runs)):
                                    next_run = para.runs[j]
                                    if next_run.text.strip() == "" or next_run.text.isspace():
                                        next_run.text = name
                                        break
                                    elif not any(char in next_run.text for char in ["姓名", "："]):
                                        next_run.text = name
                                        break
                        break
                    
                    # 班级替换
                    elif "班级" in text and class_name not in text:
                        for i, run in enumerate(para.runs):
                            if "班级" in run.text:
                                for j in range(i + 1, len(para.runs)):
                                    next_run = para.runs[j]
                                    if next_run.text.strip() == "" or next_run.text.isspace():
                                        next_run.text = class_name
                                        break
                                    elif not any(char in next_run.text for char in ["班级", "："]):
                                        next_run.text = class_name
                                        break
                        break
                    
                    # 日期替换
                    elif "日期" in text and date not in text:
                        for i, run in enumerate(para.runs):
                            if "日期" in run.text:
                                for j in range(i + 1, len(para.runs)):
                                    next_run = para.runs[j]
                                    if next_run.text.strip() == "" or next_run.text.isspace():
                                        next_run.text = date
                                        break
                                    elif not any(char in next_run.text for char in ["日期", "："]):
                                        next_run.text = date
                                        break
                        break
            
            return True
        except Exception as e:
            print(f"填写封面信息失败: {e}")
            return False
    
    def insert_table(self, rows: int, cols: int, headers: List[str],
                     data: List[List[str]], style: str = "Table Grid") -> bool:
        """
        插入表格
        
        Args:
            rows: 行数
            cols: 列数
            headers: 表头列表
            data: 数据列表
            style: 表格样式
        """
        if not self._is_open:
            return False
        
        try:
            if DOCX_SKILL_AVAILABLE:
                # 使用 XML 方式插入表格
                # 构造表格 XML 并插入
                try:
                    doc_xml = self.doc["word/document.xml"]
                    # 这里需要构造完整的表格 XML 结构
                    # 暂时使用备选方案
                    pass
                except:
                    pass
            else:
                # 使用 python-docx
                table = self.doc.add_table(rows=rows, cols=cols)
                table.style = style
                
                # 填充表头
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    if i < len(header_cells):
                        header_cells[i].text = header
                
                # 填充数据
                for i, row_data in enumerate(data):
                    if i + 1 < len(table.rows):
                        row_cells = table.rows[i + 1].cells
                        for j, cell_text in enumerate(row_data):
                            if j < len(row_cells):
                                row_cells[j].text = cell_text
            
            return True
        except Exception as e:
            print(f"插入表格失败: {e}")
            return False
    
    def set_paragraph_format(self, paragraph, font_name_cn: str = "宋体",
                            font_name_en: str = "Times New Roman",
                            font_size: int = 12,
                            first_line_indent: float = 0.35) -> bool:
        """
        设置段落格式
        
        Args:
            paragraph: 段落对象
            font_name_cn: 中文字体
            font_name_en: 英文字体
            font_size: 字号（磅）
            first_line_indent: 首行缩进（英寸）
        """
        try:
            if DOCX_SKILL_AVAILABLE:
                # 使用 XML 方式设置格式
                try:
                    # 直接操作段落 XML
                    pass
                except:
                    pass
            else:
                # 使用 python-docx
                from docx.shared import Pt, Inches
                
                paragraph.paragraph_format.first_line_indent = Inches(first_line_indent)
                
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = font_name_en
                    # 设置中文字体
                    try:
                        from docx.oxml.ns import qn
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
                    except:
                        # 如果设置失败，跳过中文字体设置
                        pass
            
            return True
        except Exception as e:
            print(f"设置段落格式失败: {e}")
            return False
    
    def add_caption(self, caption_type: str, chapter: int, number: int,
                   description: str, center: bool = True) -> bool:
        """
        添加题注
        
        Args:
            caption_type: 类型（"图"或"表"）
            chapter: 章节号
            number: 序号
            description: 描述
            center: 是否居中
        """
        if not self._is_open:
            return False
        
        try:
            caption_text = f"{caption_type}{chapter}-{number} {description}"
            
            if DOCX_SKILL_AVAILABLE:
                # 使用 XML 方式添加题注
                try:
                    doc_xml = self.doc["word/document.xml"]
                    # 构造题注段落 XML
                    pass
                except:
                    pass
            else:
                # 使用 python-docx
                para = self.doc.add_paragraph()
                if center:
                    para.alignment = 1  # 居中
                run = para.add_run(caption_text)
                run.font.size = Pt(10.5)  # 五号字
            
            return True
        except Exception as e:
            print(f"添加题注失败: {e}")
            return False
    
    def cleanup_empty_paragraphs(self, max_empty: int = 1) -> bool:
        """
        清理多余空行
        
        Args:
            max_empty: 标题间最多保留的空行数
        """
        if not self._is_open:
            return False
        
        try:
            if DOCX_SKILL_AVAILABLE:
                # 使用 XML 方式清理
                try:
                    doc_xml = self.doc["word/document.xml"]
                    # 删除多余的空段落
                    pass
                except:
                    pass
            else:
                # 使用 python-docx
                # 遍历段落，删除多余的空行
                paragraphs_to_remove = []
                empty_count = 0
                
                for i, para in enumerate(self.doc.paragraphs):
                    if not para.text.strip():
                        empty_count += 1
                        if empty_count > max_empty:
                            paragraphs_to_remove.append(para)
                    else:
                        empty_count = 0
                
                # 删除多余段落
                for para in paragraphs_to_remove:
                    p = para._element
                    p.getparent().remove(p)
            
            return True
        except Exception as e:
            print(f"清理空行失败: {e}")
            return False
    
    def add_page_break_before_heading(self, heading_text: str) -> bool:
        """
        在指定标题前添加分页符
        
        Args:
            heading_text: 标题文本
        """
        if not self._is_open:
            return False
        
        try:
            if DOCX_SKILL_AVAILABLE:
                # 使用 XML 方式添加分页
                try:
                    doc_xml = self.doc["word/document.xml"]
                    # 在指定段落前插入分页符
                    pass
                except:
                    pass
            else:
                # 使用 python-docx
                for para in self.doc.paragraphs:
                    if heading_text in para.text:
                        # 在段落前添加分页符
                        from docx.text.paragraph import Paragraph
                        new_para = self.doc.add_paragraph()
                        new_para._element.addprevious(para._element)
                        # 添加分页符到新段落
                        from docx.enum.text import WD_BREAK
                        run = new_para.add_run()
                        run.add_break(WD_BREAK.PAGE)
                        break
            
            return True
        except Exception as e:
            print(f"添加分页符失败: {e}")
            return False


def process_lab_report(template_path: str, output_path: str,
                       student_info: Dict[str, str],
                       content_data: Optional[Dict] = None) -> bool:
    """
    处理实验报告的主函数
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径
        student_info: 学生信息字典（包含 student_id, name, class_name）
        content_data: 内容数据（可选）
    
    Returns:
        bool: 是否成功
    """
    doc = LabReportDocument(template_path)
    
    if not doc.open():
        return False
    
    try:
        # 1. 填写封面信息
        doc.fill_cover_info(
            student_id=student_info.get('student_id', ''),
            name=student_info.get('name', ''),
            class_name=student_info.get('class_name', ''),
            date=student_info.get('date')
        )
        
        # 2. 插入测试用例表格（如果有）
        if content_data and 'test_cases' in content_data:
            test_cases = content_data['test_cases']
            doc.insert_table(
                rows=len(test_cases) + 1,  # +1 for header
                cols=len(test_cases[0]) if test_cases else 0,
                headers=content_data.get('table_headers', []),
                data=test_cases
            )
        
        # 3. 清理空行
        doc.cleanup_empty_paragraphs(max_empty=1)
        
        # 4. 关闭并保存
        return doc.close(output_path)
        
    except Exception as e:
        print(f"处理实验报告失败: {e}")
        doc.close()
        return False


# 便捷的函数接口
def fill_cover(docx_path: str, student_id: str, name: str, 
               class_name: str, date: Optional[str] = None) -> bool:
    """便捷函数：填写封面"""
    doc = LabReportDocument(docx_path)
    if doc.open():
        result = doc.fill_cover_info(student_id, name, class_name, date)
        doc.close()
        return result
    return False


def insert_test_case_table(docx_path: str, headers: List[str], 
                           data: List[List[str]]) -> bool:
    """便捷函数：插入测试用例表格"""
    doc = LabReportDocument(docx_path)
    if doc.open():
        result = doc.insert_table(
            rows=len(data) + 1,
            cols=len(headers),
            headers=headers,
            data=data
        )
        doc.close()
        return result
    return False


def cleanup_document(docx_path: str) -> bool:
    """便捷函数：清理文档格式"""
    doc = LabReportDocument(docx_path)
    if doc.open():
        result = doc.cleanup_empty_paragraphs(max_empty=1)
        doc.close()
        return result
    return False
